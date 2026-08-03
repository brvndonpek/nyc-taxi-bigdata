"""
Generates report/IST3134_Taxi_BigData_Report.docx from the analysis outputs.
Run after the analysis + charts:  python src/make_report.py
"""
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "output"
FIG = "output/figures"
BLUE = RGBColor(0x1f, 0x4e, 0x79)


def hr(doc):
    doc.add_paragraph()


def add_table(doc, df, max_rows=None, colwidths=None):
    d = df if max_rows is None else df.head(max_rows)
    t = doc.add_table(rows=1, cols=len(d.columns))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(d.columns):
        run = t.rows[0].cells[j].paragraphs[0].add_run(str(c))
        run.bold = True
        run.font.size = Pt(9)
    for _, row in d.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(d.columns):
            v = row[c]
            if isinstance(v, float):
                v = f"{v:,.2f}" if abs(v) < 1e6 else f"{v:,.0f}"
            elif isinstance(v, (int,)):
                v = f"{v:,}"
            p = cells[j].paragraphs[0]
            p.add_run(str(v)).font.size = Pt(9)
    return t


def fig(doc, path, caption, width=6.2):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE
    return p


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    # -------- Title page --------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Big Data Analytics of NYC Yellow Taxi Trips (2024)")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = BLUE
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A PySpark vs. Pandas Comparison of Demand and Tipping Patterns")
    r.font.size = Pt(13)
    r.italic = True
    for _ in range(2):
        doc.add_paragraph()
    for line in [
        "IST3134 — Big Data Analytics",
        "Group Assignment, May Semester 2026",
        "",
        "Member 1:  ____________________  (Student ID: __________)",
        "Member 2:  ____________________  (Student ID: __________)",
        "",
        "Source code & dataset: https://github.com/brvndonpek/nyc-taxi-bigdata",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(12)
    doc.add_page_break()

    # -------- 1. Problem --------
    h(doc, "1. Introduction to the Problem")
    doc.add_paragraph(
        "New York City's yellow taxis complete tens of millions of trips every "
        "year. Understanding when and where demand concentrates — and how riders "
        "tip — is valuable to several stakeholders: the Taxi & Limousine "
        "Commission (fleet regulation and driver supply), taxi operators "
        "(positioning cars to cut idle time), and drivers themselves (choosing "
        "shifts and locations that maximise earnings). The core questions this "
        "project answers are:")
    for b in [
        "How does trip demand vary across the 24 hours of the day?",
        "Which pickup zones and boroughs generate the most trips and revenue?",
        "How does tipping behaviour (tip as a % of fare) change by time of day, "
        "location, and payment method?",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(
        "The challenge is one of scale. A single year of trip records is over 41 "
        "million rows (~660 MB of compressed Parquet). Answering these questions "
        "requires aggregating across the entire dataset repeatedly. This makes it "
        "a natural fit for a distributed, big-data processing approach, and a good "
        "case study for comparing a big-data engine (Apache Spark) against a "
        "conventional single-machine tool (Pandas).")

    # -------- 2. Dataset --------
    h(doc, "2. Introduction to the Dataset")
    doc.add_paragraph(
        "The data is the publicly available Yellow Taxi Trip Records published by "
        "the New York City Taxi & Limousine Commission (TLC). We use all twelve "
        "monthly files for 2024.")
    for b in [
        "Provider: NYC Taxi & Limousine Commission (TLC).",
        "Landing page: https://www.nyc.gov/site/tlc/about/tlc-trip-record-data.page",
        "Files: yellow_tripdata_2024-01.parquet … 2024-12.parquet (Apache Parquet).",
        "Volume: 41,169,720 trips; ~660 MB compressed.",
        "Zone lookup: taxi_zone_lookup.csv maps LocationID to borough and zone name.",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(
        "Each row is one completed trip. The fields used in this analysis are:")
    schema = pd.DataFrame({
        "Field": ["tpep_pickup_datetime", "tpep_dropoff_datetime",
                  "trip_distance", "PULocationID", "payment_type",
                  "fare_amount", "tip_amount", "total_amount",
                  "passenger_count"],
        "Meaning": ["Trip start timestamp", "Trip end timestamp",
                    "Distance in miles", "Pickup zone ID (join to lookup)",
                    "1=card, 2=cash, …", "Metered fare (US$)",
                    "Tip amount (US$)", "Total charged (US$)",
                    "Number of passengers"],
    })
    add_table(doc, schema)
    doc.add_paragraph(
        "Data quality note: the raw feed contains erroneous records (zero or "
        "negative fares, zero-distance trips, impossible distances > 100 miles, "
        "zero-passenger trips, and a few timestamps outside 2024). We filter "
        "these out before analysis, retaining 35,628,429 trips — 86.5% of the raw "
        "data. Applying identical filters in both implementations keeps the "
        "comparison fair and the results trustworthy.")

    # -------- 3. Approach --------
    h(doc, "3. The MapReduce / Spark Approach")
    doc.add_paragraph(
        "The analysis is fundamentally a set of group-by aggregations, which is "
        "exactly the problem the MapReduce paradigm was designed for. Taking "
        "“average tip % and trip count per pickup zone” as the example:")
    doc.add_paragraph(
        "Map phase: for every trip, emit a key-value pair whose key is the "
        "grouping attribute (e.g. the pickup zone, or the hour of day) and whose "
        "value carries the measures we need — the fare, tip, distance and a count "
        "of 1.", style="List Bullet")
    doc.add_paragraph(
        "Shuffle phase: the framework groups all values that share the same key "
        "onto the same reducer. This is the step that a single machine cannot "
        "parallelise but a cluster can.", style="List Bullet")
    doc.add_paragraph(
        "Reduce phase: for each key, sum the counts and the measures, then derive "
        "the averages (total tip% ÷ count, etc.).", style="List Bullet")
    doc.add_paragraph(
        "We implement this with Apache Spark's DataFrame API in PySpark. Although "
        "we write concise groupBy().agg() calls, Spark compiles them into exactly "
        "this map-shuffle-reduce plan and executes it across all CPU cores (and, "
        "on a cluster, across many machines). Spark reads the Parquet files as a "
        "distributed DataFrame, partitions the 41M rows, runs the map and partial "
        "aggregation on each partition in parallel, shuffles by key, and finishes "
        "the reduction. Key implementation choices:")
    for b in [
        "Parquet columnar format lets Spark read only the columns each query "
        "needs (predicate/column pruning), cutting I/O dramatically.",
        "The small zone-lookup table is broadcast to every worker, so the "
        "join needs no shuffle of the big table.",
        "The cleaned DataFrame is cached because it is reused by four separate "
        "aggregations, avoiding a re-read of the data each time.",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(
        "For comparison we implement the identical logic in Pandas on a single "
        "machine. Because 41M rows with derived columns will not comfortably fit "
        "in 16 GB of RAM, the Pandas version must be engineered carefully: it "
        "reads only the required columns and processes the twelve files one at a "
        "time, combining partial aggregates at the end. Spark performs this "
        "out-of-core, parallel execution automatically — the extra manual effort "
        "on the Pandas side is itself part of the finding.")

    h(doc, "3.1 A worked example", level=2)
    doc.add_paragraph(
        "To make the map-shuffle-reduce flow concrete, consider computing the "
        "average tip % per hour from five sample trips. In the map phase each "
        "trip is turned into a (key, value) pair keyed on the pickup hour, "
        "carrying the tip % and a count of 1:")
    ex_map = pd.DataFrame({
        "Trip": ["A", "B", "C", "D", "E"],
        "Pickup hour": [18, 9, 18, 9, 18],
        "tip %": [20.0, 15.0, 30.0, 25.0, 22.0],
        "Map output (key → value)": [
            "18 → (20.0, 1)", "9 → (15.0, 1)", "18 → (30.0, 1)",
            "9 → (25.0, 1)", "18 → (22.0, 1)"],
    })
    add_table(doc, ex_map)
    doc.add_paragraph(
        "The shuffle phase then groups all values that share a key onto the same "
        "reducer — the trips for hour 18 land together, and likewise for hour 9. "
        "Finally the reduce phase sums the tip percentages and the counts per "
        "key and divides to get the average:")
    ex_red = pd.DataFrame({
        "Hour (key)": [9, 18],
        "Values after shuffle": ["(15.0,1),(25.0,1)", "(20.0,1),(30.0,1),(22.0,1)"],
        "Reduce: sum tip% / sum count": ["40.0 / 2", "72.0 / 3"],
        "Avg tip %": [20.0, 24.0],
    })
    add_table(doc, ex_red)
    doc.add_paragraph(
        "This is exactly what runs at scale in our Spark job — only instead of "
        "five rows on one machine, the 35.6M cleaned trips are split into "
        "partitions, mapped in parallel across all cores, shuffled by key, and "
        "reduced. Because each partition is mapped independently, the work scales "
        "almost linearly with the number of cores or machines available — the "
        "property that makes the approach a “big data” solution.")

    # -------- 4. Analysis of output --------
    h(doc, "4. Analysis of the Output")

    h(doc, "4.1 Demand and tipping by hour of day", level=2)
    fig(doc, f"{FIG}/fig1_demand_by_hour.png",
        "Figure 1. Trip volume (bars) and average tip % (line) by hour of day.")
    doc.add_paragraph(
        "Demand follows a clear daily cycle. It bottoms out at 4–5 am (~170k–200k "
        "trips) and rises through the day to a peak at 6 pm (2.56M trips) — the "
        "evening rush. Tipping is highest in the evening (peaking around 23–24% "
        "between 6–8 pm) and lowest during the pre-dawn commute hours. The one "
        "anomaly, a tip-% spike at 4 am, comes from the very small number of "
        "trips at that hour, where a few generous tips move the average — a good "
        "reminder to read averages alongside their counts.")

    h(doc, "4.2 Busiest pickup zones", level=2)
    tz = pd.read_csv(f"{OUT}/spark_top_zones.csv").head(10)
    tz = tz[["Borough", "Zone", "trips", "avg_fare", "avg_tip_pct"]]
    add_table(doc, tz)
    fig(doc, f"{FIG}/fig2_top_zones.png",
        "Figure 2. Top 10 pickup zones by trip count.")
    doc.add_paragraph(
        "JFK Airport is the single busiest pickup zone (1.86M trips) and, with an "
        "average fare of about $64, by far the most lucrative per trip — these are "
        "long airport runs. LaGuardia Airport is similar ($44 average fare). The "
        "remaining top zones are dense Manhattan business and nightlife districts "
        "(Upper East Side, Midtown, Penn Station, Times Square), which generate "
        "huge trip volumes at lower per-trip fares. Notably, airport riders tip a "
        "lower percentage (JFK ~14.6%) than Manhattan riders (~22–23%).")

    h(doc, "4.3 Revenue and tipping by borough", level=2)
    bo = pd.read_csv(f"{OUT}/spark_by_borough.csv")
    bo = bo[bo.trips > 1000][["Borough", "trips", "total_revenue", "avg_tip_pct"]]
    add_table(doc, bo)
    fig(doc, f"{FIG}/fig3_revenue_by_borough.png",
        "Figure 3. Total revenue by pickup borough.")
    doc.add_paragraph(
        "Manhattan overwhelmingly dominates yellow-taxi activity, accounting for "
        "roughly $757M of pickup revenue — close to 89% of the total — followed by "
        "Queens (driven almost entirely by the two airports). The outer boroughs "
        "(Brooklyn, Bronx, Staten Island) see very little yellow-taxi pickup "
        "activity, consistent with the fact that they are largely served by green "
        "taxis and for-hire vehicles. This concentration is a useful signal for "
        "fleet positioning.")

    h(doc, "4.4 Tipping by payment method — a data-quality insight", level=2)
    pay = pd.read_csv(f"{OUT}/spark_by_payment.csv")
    names = {1: "Credit card", 2: "Cash", 3: "No charge", 4: "Dispute"}
    pay = pay[pay.payment_type.isin(names)].copy()
    pay["Payment method"] = pay.payment_type.map(names)
    add_table(doc, pay[["Payment method", "trips", "avg_tip_pct"]])
    fig(doc, f"{FIG}/fig5_payment.png",
        "Figure 5. Average tip % by payment method.")
    doc.add_paragraph(
        "This breakdown exposes an important characteristic of the data. "
        "Credit-card trips (29.9M, the large majority) average a 25.1% tip, but "
        "cash trips (5.2M) record an average tip of 0.0%. Cash passengers "
        "obviously do tip — the explanation is that the TLC data set only "
        "captures tips entered through the card reader; cash tips are never "
        "recorded. This matters for interpretation: the ~21% blended tip average "
        "reported earlier is diluted by cash trips showing as zero, so the true "
        "card-tipping rate (25.1%) is the more meaningful figure. Recognising "
        "this prevents a wrong conclusion and demonstrates why understanding how "
        "data is collected is as important as the analysis itself.")

    h(doc, "4.5 Engine comparison: PySpark vs Pandas", level=2)
    bench = pd.read_csv(f"{OUT}/benchmark.csv")
    show = bench.rename(columns={
        "months": "Months", "rows": "Rows",
        "pandas_sec": "Pandas (s)", "pandas_peak_mb": "Pandas peak MB",
        "spark_sec": "Spark (s)"})
    add_table(doc, show)
    fig(doc, f"{FIG}/fig4_benchmark.png",
        "Figure 4. Runtime and Pandas memory as the dataset grows.")
    doc.add_paragraph(
        "Both engines produce identical numerical results, confirming the "
        "implementations are equivalent. The interesting story is performance. On "
        "a single month Pandas is much faster (0.2 s vs 3.4 s) because Spark pays "
        "a fixed start-up cost for its JVM, query planner and shuffle machinery. "
        "But as data grows the picture inverts: by the full year Spark (1.0 s) is "
        "three times faster than Pandas (3.0 s), and — more importantly — Pandas' "
        "peak memory rises linearly with data, from 212 MB to nearly 4 GB. "
        "Extrapolating, at roughly five times this volume Pandas would exhaust a "
        "16 GB machine and fail, whereas Spark's runtime stays almost flat and it "
        "can spill to disk and scale out across a cluster.")
    doc.add_paragraph(
        "The conclusion is nuanced and is the central lesson of the assignment: a "
        "big-data engine is not automatically faster — for data that fits in "
        "memory, a single-machine tool can win. Spark's value is scalability and "
        "robustness: it handles datasets larger than RAM with simple code and "
        "grows to cluster scale, which is precisely what “big data” "
        "demands. Pandas only kept up here because we manually chunked the input; "
        "Spark required no such effort.")

    # -------- 5. Reflection --------
    h(doc, "5. Individual Reflection")
    note = doc.add_paragraph()
    r = note.add_run(
        "NOTE TO AUTHORS: the two reflections below are drafts grounded in the "
        "real work done in this project. Personalise them in your own words and "
        "adjust any detail (e.g. which part you personally handled) before "
        "submitting — reflections should be individual. Delete this note.")
    r.italic = True

    h(doc, "Member 1 — <name>", level=2)
    doc.add_paragraph(
        "The biggest thing I took away from this assignment is that “big data” is "
        "about scalability, not raw speed. Before starting, I assumed Spark would "
        "always beat Pandas. Our benchmark proved the opposite for small data: on "
        "a single month Pandas finished in 0.2 s while Spark took 3.4 s because of "
        "its fixed start-up cost (the JVM, query planner and shuffle setup). Only "
        "as the data grew to the full 35.6M rows did Spark pull clearly ahead "
        "(1.0 s vs 3.0 s). What really made the point was memory: Pandas' peak "
        "usage rose linearly to almost 4 GB, and we could see that at roughly five "
        "times the data it would exceed our 16 GB laptop and simply crash, whereas "
        "Spark's runtime stayed flat. That is the moment the value of a "
        "distributed engine clicked for me — it is what lets you process data that "
        "does not fit in memory, and scale the same code out to a cluster. "
        "Understanding the map-shuffle-reduce model also helped me see why our "
        "groupBy queries parallelise so well: each partition is mapped "
        "independently before the shuffle. If I did this again I would run it on a "
        "real cloud cluster (e.g. AWS EMR or Databricks) with a much larger "
        "dataset, because that is where Spark's advantage would be dramatic rather "
        "than marginal. (Approx. 230 words — edit to your own experience.)")

    h(doc, "Member 2 — <name>", level=2)
    doc.add_paragraph(
        "My main lesson was that understanding the data matters as much as the "
        "processing technology. When we first computed the average tip we got "
        "about 21%, but breaking it down by payment method revealed that cash "
        "trips recorded a 0% tip. Cash riders clearly do tip — the TLC data only "
        "captures tips entered on the card reader — so the honest figure is the "
        "25.1% card-tipping rate. Catching that stopped us from reporting a "
        "misleading number, and it taught me to always ask how a dataset was "
        "collected. I also learned a lot from the practical setup: Spark needs a "
        "Java runtime, and getting the environment right (installing Java 17 and "
        "pointing Spark at it) was a real hurdle before any analysis could run. "
        "Data cleaning was another eye-opener — 13.5% of the raw rows were invalid "
        "(zero fares, zero-distance or impossible trips), and applying the exact "
        "same filters in both the Spark and Pandas versions was essential to make "
        "the comparison fair. Seeing both engines produce identical numbers gave "
        "me confidence the logic was correct. Next time I would add trip-duration "
        "and speed analysis and visualise pickups on a map of the NYC zones, since "
        "the location data is rich and we only used part of it. "
        "(Approx. 230 words — edit to your own experience.)")

    # -------- Appendix --------
    h(doc, "Appendix — Reproducing the Results")
    doc.add_paragraph(
        "All source code and instructions are in the GitHub repository linked on "
        "the title page. The dataset is downloaded from the TLC site (link in "
        "Section 2 and the README). To reproduce every table and figure:")
    for c in ["pip install -r requirements.txt",
              "conda install -c conda-forge openjdk=17   # Java 17 for Spark",
              "bash run_all.sh"]:
        p = doc.add_paragraph()
        p.add_run(c).font.name = "Consolas"
        p.runs[0].font.size = Pt(10)

    doc.save("report/IST3134_Taxi_BigData_Report.docx")
    print("Saved report/IST3134_Taxi_BigData_Report.docx")


if __name__ == "__main__":
    main()
